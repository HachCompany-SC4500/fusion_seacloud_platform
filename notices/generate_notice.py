#!/usr/bin/env python3

import os
import json
from collections import OrderedDict

# For docx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# For PDF
from fpdf import FPDF

from abc import ABC, abstractmethod

CONTROLLER_NAME = 'SC4500'
CONTROLLER_BEFE = 'Backend_Frontend'
DEPLOY_LOCATION = '../../../deploy'
DEPLOY_IMAGE_LOCATION = '../../../deploy/images/colibri-imx7-emmc-1370'
LICENSE_LOCATION = os.path.join(DEPLOY_LOCATION, 'licenses')
BEFE_LICENSE_LOCATION = os.path.join(DEPLOY_IMAGE_LOCATION, 'BEFE-licenses.json')

class DocumentGenerator(ABC):
    """
    Abstract class for DocumentGenerators
    """
    def __init__(self):
        super().__init__()

    @abstractmethod
    def add_title(self, text):
        """Add a title"""
        pass

    @abstractmethod
    def add_heading(self, text, level):
        """Add a heading"""
        pass

    @abstractmethod
    def add_paragraph(self, text, font_size = None):
        """Add a paragraph"""
        pass

    @abstractmethod
    def add_page_break(self):
        """ Add a page break"""
        pass

    @abstractmethod
    def save(self, filename):
        """
        Save the file
        @filename : doesn't include extension
        """
        pass

class DocxGenerator(DocumentGenerator):
    """ Generator for docx document"""
    def __init__(self):
        super().__init__()
        self.document = Document()

    def add_title(self, text):
        self.document.add_heading(text, level = 0).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_heading(self, text, level):
        self.document.add_heading(text, level)

    def add_paragraph(self, text, font_size = None):
        if not font_size:
            self.document.add_paragraph(text)
        else:
            run = self.document.add_paragraph().add_run(text)
            font = run.font
            font.size = Pt(font_size)

    def add_page_break(self):
        self.document.add_page_break()

    def save(self, filename):
        self.document.save(filename + '.docx')

class PdfGenerator(DocumentGenerator):
    """Generator for PDF document"""
    def __init__(self):
        super().__init__()
        self.pdf = FPDF()
        self.pdf.set_margins(31.7, 25.4)
        self.pdf.add_page()
        self.default_fontname = 'Arial'
        self.default_fontsize = 11
        self.set_default_font()

    def set_default_font(self):
        self.pdf.set_font(self.default_fontname, '', self.default_fontsize)

    def add_title(self, text):
        self.pdf.set_font(self.default_fontname, '', 26)
        self.pdf.set_text_color(23, 54, 93)
        self.pdf.multi_cell(0, 10, text, border = 'B', align = 'C', fill = False)
        self.pdf.ln()
        self.pdf.set_text_color(0,0,0)
        self.set_default_font()

    def add_heading(self, text, level):
        if level <= 1:
            self.pdf.set_text_color(54, 95, 145)
            self.pdf.set_font(self.default_fontname, 'B')
            self.add_paragraph(text, font_size = 14)
        elif level >= 2:
            self.pdf.set_text_color(79, 129, 189)
            self.pdf.set_font(self.default_fontname, 'B')
            self.add_paragraph(text, font_size = 13)
        self.pdf.set_text_color(0,0,0)
        self.set_default_font()

    def add_paragraph(self, text, font_size = None):
        if font_size is None:
            font_size = self.default_fontsize
        self.pdf.set_font_size(font_size)
        self.pdf.write(font_size/2, txt = text)
        self.pdf.ln()

    def add_page_break(self):
        self.pdf.add_page()

    def save(self, filename):
        self.pdf.output(filename + '.pdf', 'F')

print("Generate Notice summary file for {} controller".format(CONTROLLER_NAME))
print("Using information found in folder {}".format(LICENSE_LOCATION))

generators = [PdfGenerator()]
#generators.append(DocxGenerator())

# Generate licenses file for OS part
for document in generators:
    document.add_title("{} Open Source Software Notices and Licenses".format(CONTROLLER_NAME))

    document.add_paragraph('This file contains attributions, copyright notices and licenses associated with Open Source Software used in {} firmware.'.format(CONTROLLER_NAME))

    document.add_page_break()

    for dirpath, dirnames, files in os.walk(LICENSE_LOCATION):

        # Keep only current folder name
        recipe_name = dirpath.split(os.sep)[-1]

        # Sort subfolders
        dirnames.sort()

        recipeinfo_path = os.path.join(dirpath, 'recipeinfo')

        # Skip folder without recipes
        if not os.path.isfile(recipeinfo_path):
            continue

        # Skip native recipes that are not part of the image published
        if '-native' in recipe_name:
            continue

        document.add_heading(recipe_name, level=1)

        # Extract package information from recipeinfo file generated by Yocto
        package_info = {}
        with open(recipeinfo_path) as recipe_info:
            for line in recipe_info.readlines():
                fields = line.split(':')
                package_info[fields[0]] = fields[1].strip()

        document.add_paragraph('Package version: {}\nPackage revision: {}\nPackage license(s): {}\n'.format(package_info['PV'], package_info['PR'], package_info['LICENSE']))

        # Print license related files attached by Yocto to the package (license files, copying, readme etc)
        for file in files:
            # Filter recipeinfo file and other unexpected source files
            if 'recipeinfo' in file:
                continue
            document.add_heading('File name: {}'.format(file), level=2)
            with open(os.path.join(dirpath, file), encoding='iso-8859-1') as content:
                file_data = content.read()
                # Remove XML standard unsupported control characters (currently only Form Feed character (0x0C) cause a problem)
                file_data = ''.join(c for c in file_data if ord(c) >= 160 or (ord(c) >= 32 and ord(c) <= 128 ) or ord(c) == 0x09 or ord(c) == 0x0A or ord(c) == 0x0D)
                document.add_paragraph('{}'.format(file_data), font_size=6)

        document.add_page_break()

    document.save(os.path.join(DEPLOY_LOCATION, '{}_OSS_Notices'.format(CONTROLLER_NAME)))

# generate pdf file for backend / frontend licenses
print("Generate Notice summary file for {} {}".format(CONTROLLER_NAME,CONTROLLER_BEFE))
print("Using information found in {}".format(BEFE_LICENSE_LOCATION))

befe_document = PdfGenerator()
befe_document.add_title("{} {} Open Source Software Notices and Licences".format(CONTROLLER_NAME,CONTROLLER_BEFE))

befe_document.add_paragraph('This file contains attributions, copyright notices and licenses associated with Open Source Software used in {} {} firmware.'.format(CONTROLLER_NAME,CONTROLLER_BEFE))

with open(BEFE_LICENSE_LOCATION,encoding='iso-8859-1') as json_file:
    data = json.load(json_file, object_pairs_hook=OrderedDict)

    for node in data:
        befe_document.add_page_break()
        befe_document.add_heading(node,1)
        # populate information with expected order
        if 'licenses' in  data[node]:
            befe_document.add_heading('License',2)
            befe_document.add_paragraph(data[node]['licenses'])
        if 'publisher' in  data[node]:
            befe_document.add_heading('Publisher',2)
            befe_document.add_paragraph(data[node]['publisher'])
        if 'repository' in  data[node]:
            befe_document.add_heading('Repository',2)
            befe_document.add_paragraph(data[node]['repository'])
        if 'licenseFile' in  data[node]:
            befe_document.add_heading('License file',2)
            befe_document.add_paragraph(data[node]['licenseFile'], font_size=6)

befe_document.save(os.path.join(DEPLOY_LOCATION, '{}_{}_OSS_Notices'.format(CONTROLLER_NAME,CONTROLLER_BEFE))) 

